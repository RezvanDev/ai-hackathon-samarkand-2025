import os
import pytesseract
from PyPDF2 import PdfReader
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from docx import Document
import pandas as pd
import logging
from .google_docs_parser import parse_google_doc

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def save_attachment(part, download_folder):
    filename = part.get_filename()
    if filename:
        filepath = os.path.join(download_folder, filename)
        os.makedirs(download_folder, exist_ok=True)
        with open(filepath, 'wb') as f:
            f.write(part.get_payload(decode=True))
        return filepath
    return None


def parse_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"Страница {page_num + 1}:\n{page_text}")
            except Exception as e:
                logger.warning(f"Ошибка при обработке страницы {page_num + 1}: {e}")
                continue

        if not text_parts:
            logger.info("Текст не извлечен, пробуем OCR для PDF")
            return parse_pdf_with_ocr(filepath)

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка при парсинге PDF {filepath}: {e}")
        raise ValueError(f"Не удалось обработать PDF файл: {e}")


def parse_pdf_with_ocr(filepath):
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(filepath)
        text_parts = []

        for i, image in enumerate(images):
            img = image.convert('L')
            img = ImageOps.autocontrast(img)

            configs = [
                '--psm 6',
                '--psm 3',
                '--psm 1'
            ]

            for config in configs:
                try:
                    text = pytesseract.image_to_string(img, config=config, lang='rus+eng')
                    if text.strip():
                        text_parts.append(f"Страница {i + 1}:\n{text}")
                        break
                except Exception as e:
                    logger.warning(f"OCR с конфигом {config} не удался: {e}")
                    continue

        return "\n\n".join(text_parts) if text_parts else "Текст не удалось извлечь"
    except ImportError:
        logger.warning("pdf2image не установлен, OCR недоступен")
        return "Текст не удалось извлечь (требуется pdf2image)"
    except Exception as e:
        logger.error(f"Ошибка OCR для PDF {filepath}: {e}")
        return "Текст не удалось извлечь"

def parse_image(filepath):
    try:
        img = Image.open(filepath)
        methods = [
            lambda img: img.convert('L'),
            lambda img: img.convert('L').filter(ImageFilter.SHARPEN),
            lambda img: ImageOps.autocontrast(img.convert('L')),
            lambda img: ImageEnhance.Contrast(img.convert('L')).enhance(2.0),
            lambda img: ImageEnhance.Sharpness(img.convert('L')).enhance(2.0)
        ]

        configs = [
            '--psm 6',
            '--psm 3',
            '--psm 1',
            '--psm 4',
            '--psm 8'
        ]

        best_result = ""

        for method in methods:
            try:
                processed_img = method(img)

                for config in configs:
                    try:
                        text = pytesseract.image_to_string(processed_img, config=config, lang='rus+eng')
                        if len(text.strip()) > len(best_result):
                            best_result = text.strip()
                    except Exception as e:
                        logger.warning(f"OCR с конфигом {config} не удался: {e}")
                        continue

            except Exception as e:
                logger.warning(f"Метод обработки изображения не удался: {e}")
                continue

        if not best_result:
            try:
                best_result = pytesseract.image_to_string(img, lang='rus+eng')
            except Exception as e:
                logger.error(f"OCR не удался для {filepath}: {e}")
                return "Текст не удалось извлечь из изображения"

        return best_result
    except Exception as e:
        logger.error(f"Ошибка при обработке изображения {filepath}: {e}")
        raise ValueError(f"Не удалось обработать изображение: {e}")


def parse_doc(filepath):
    try:
        doc = Document(filepath)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка при парсинге Word документа {filepath}: {e}")
        raise ValueError(f"Не удалось обработать Word документ: {e}")

def parse_excel(filepath):
    try:
        engines = ['openpyxl', 'xlrd']

        for engine in engines:
            try:
                if engine == 'xlrd':
                    data = pd.read_excel(filepath, engine=engine)
                else:
                    data = pd.read_excel(filepath, engine=engine)
                result = []

                if not data.empty:
                    headers = data.columns.tolist()
                    result.append(" | ".join(str(h) for h in headers))
                    for _, row in data.iterrows():
                        row_data = [str(val) for val in row.values]
                        result.append(" | ".join(row_data))

                return "\n".join(result)

            except Exception as e:
                logger.warning(f"Движок {engine} не удался: {e}")
                continue

        raise ValueError("Не удалось обработать Excel файл ни одним из движков")

    except Exception as e:
        logger.error(f"Ошибка при парсинге Excel файла {filepath}: {e}")
        raise ValueError(f"Не удалось обработать Excel файл: {e}")


def parse_text_file(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Ошибка при чтении текстового файла {filepath}: {e}")
            raise ValueError(f"Не удалось прочитать текстовый файл: {e}")

def parse_attachment(filepath):
    if not os.path.exists(filepath):
        raise ValueError(f"Файл не найден: {filepath}")

    filename = os.path.basename(filepath)
    ext = filename.split('.')[-1].lower() if '.' in filename else ''

    logger.info(f"Обработка файла: {filename} (расширение: {ext})")

    try:
        if ext == "pdf":
            return parse_pdf(filepath)
        elif ext in ["jpg", "jpeg", "png", "bmp", "tiff", "tif"]:
            return parse_image(filepath)
        elif ext in ["docx", "doc"]:
            return parse_doc(filepath)
        elif ext in ["xlsx", "xls"]:
            return parse_excel(filepath)
        elif ext in ["txt", "csv"]:
            return parse_text_file(filepath)
        else:
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(4)

                if header.startswith(b'%PDF'):
                    logger.info("Обнаружен PDF по магическому числу")
                    return parse_pdf(filepath)
                elif header.startswith(b'\xff\xd8\xff'):
                    logger.info("Обнаружено изображение JPEG по магическому числу")
                    return parse_image(filepath)
                elif header.startswith(b'\x89PNG'):
                    logger.info("Обнаружено изображение PNG по магическому числу")
                    return parse_image(filepath)
                else:
                    logger.info("Пробуем обработать как текстовый файл")
                    return parse_text_file(filepath)

            except Exception as e:
                logger.error(f"Не удалось определить тип файла {filepath}: {e}")
                raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    except Exception as e:
        logger.error(f"Ошибка при обработке файла {filepath}: {e}")
        raise ValueError(f"Ошибка обработки файла: {e}")
